import re
from datetime import date, datetime, timedelta
from pathlib import Path

import docx
from django.core.management.base import BaseCommand, CommandError

from myapp.models import Menu


# NEW format (April 2026+): each block has explicit day name in row 0,
# 'Menu'/'Ingredients' subheaders in row 1.
#   row 0       → day name header (e.g., "Monday")
#   row 1       → "Menu" / "Ingredients" sub-headers
#   row 2       → 8:00  (cold breakfast)
#   row 4       → 10:00 (hot breakfast)
#   row 6       → 12:00 (lunch)
#   row 11      → 5:00  (dinner)
BLOCK_SIZE = 12
MEAL_ROWS = {
    2:  'cold_breakfast',
    4:  'hot_breakfast',
    6:  'lunch',
    11: 'dinner',
}
DAY_OFFSET = {
    'monday':    0,
    'tuesday':   1,
    'wednesday': 2,
    'thursday':  3,
    'friday':    4,
}

# OLD format (Jan-Mar 2026): row 0 = "Week of M/D-M/D", days are implicit
# by block index, two days per block in cols 2-3 and 5-6.
#   row 0       → "Week of 1/12-1/18" header
#   block 0 (rows 2-13): Mon/Tue
#   block 1 (rows 14-25): Wed/Thu
#   block 2 (rows 26-37): Fri/Sat
# Within each block:
#   offset 0    → "Itinerary"/"Notes" sub-header
#   offset 1    → cold_breakfast
#   offset 3    → hot_breakfast
#   offset 5    → lunch
#   offset 10   → dinner
OLD_BLOCK_SIZE = 12
OLD_MEAL_ROWS = {
    1:  'cold_breakfast',
    3:  'hot_breakfast',
    5:  'lunch',
    10: 'dinner',
}
OLD_MENU_COLS = [2, 5]   # menu text columns within each block
OLD_DAY_PAIRS = [
    ('monday',    'tuesday'),
    ('wednesday', 'thursday'),
    ('friday',    'saturday'),
]
MEAL_PREFIX_RE = re.compile(
    r'^\s*(cold breakfast|hot breakfast|lunch|dinner)\s*:\s*(.*)$',
    re.IGNORECASE,
)
FILENAME_RANGE_RE = re.compile(r'(\d{1,2})\.(\d{1,2})\s*-\s*(\d{1,2})\.(\d{1,2})')


def parse_date_range(filename: str, year: int) -> tuple[date, date]:
    """Extract (week1_start, week2_start) from 'Kitchen Coordination 4.13 - 4.26.docx'."""
    m = FILENAME_RANGE_RE.search(filename)
    if not m:
        raise CommandError(f"Could not parse date range from filename: {filename}")
    m1, d1, _, _ = (int(x) for x in m.groups())
    start = date(year, m1, d1)
    return start, start + timedelta(days=7)


def strip_meal_prefix(text: str, fallback_slot: str) -> tuple[str, str]:
    """Return (slot, dish). If text has 'Cold Breakfast: X', slot is inferred; else use fallback."""
    m = MEAL_PREFIX_RE.match(text)
    if m:
        slot_text = m.group(1).lower().replace(' ', '_')
        return slot_text, m.group(2).strip()
    return fallback_slot, text.strip()


def find_menu_columns(subheader_cells: list[str]) -> list[int]:
    """Indices of cells containing 'Menu' (case-insensitive)."""
    return [i for i, c in enumerate(subheader_cells) if c.strip().lower() == 'menu']


def cell_text(cells, idx: int) -> str:
    if idx >= len(cells):
        return ''
    return cells[idx].text.strip()


def detect_format(table) -> str:
    """Return 'new' or 'old' based on row 0 content."""
    if not table.rows:
        return 'new'
    row0_text = ' '.join(c.text for c in table.rows[0].cells).lower()
    if 'week of' in row0_text:
        return 'old'
    return 'new'


def parse_block_old(table, block_start: int, block_idx: int, week_start: date) -> list[dict]:
    """Parse one OLD-format block. block_idx 0/1/2 → Mon-Tue / Wed-Thu / Fri-Sat."""
    if block_idx >= len(OLD_DAY_PAIRS):
        return []
    day_pair = OLD_DAY_PAIRS[block_idx]
    out = []
    for col_idx, day_name in zip(OLD_MENU_COLS, day_pair):
        if not day_name or day_name not in DAY_OFFSET:
            continue
        day_date = week_start + timedelta(days=DAY_OFFSET[day_name])
        for row_offset, default_slot in OLD_MEAL_ROWS.items():
            row_n = block_start + row_offset
            if row_n >= len(table.rows):
                continue
            row = table.rows[row_n].cells
            menu_text = cell_text(row, col_idx)
            if not menu_text:
                continue
            slot, dish = strip_meal_prefix(menu_text, default_slot)
            if not dish or dish.lower() in ('n/a', 'none', ''):
                continue
            # OLD format has Notes (cook name) in col_idx+1; keep as ingredients
            # so it shows up somewhere — these old docs didn't list ingredients.
            ingredients = cell_text(row, col_idx + 1)
            out.append({
                'date':            day_date,
                'meal_slot':       slot,
                'dish_freetext':   dish[:200],
                'ingredients_raw': ingredients,
            })
    return out


def parse_block(table, block_start: int, week_start: date) -> list[dict]:
    """Parse one 12-row block; return a list of Menu field dicts (Mon-Fri only)."""
    header_cells  = table.rows[block_start].cells
    subhdr_cells  = table.rows[block_start + 1].cells

    menu_cols = find_menu_columns([c.text for c in subhdr_cells])
    out = []

    for menu_col in menu_cols:
        day_name = header_cells[menu_col].text.strip().lower()
        if day_name not in DAY_OFFSET:
            continue  # skip Saturday/Sunday and blanks
        day_date = week_start + timedelta(days=DAY_OFFSET[day_name])
        ing_col  = menu_col + 1

        for row_offset, default_slot in MEAL_ROWS.items():
            row = table.rows[block_start + row_offset].cells
            menu_text = cell_text(row, menu_col)
            if not menu_text:
                continue
            slot, dish = strip_meal_prefix(menu_text, default_slot)
            if not dish or dish.lower() in ('n/a', 'none', ''):
                continue
            ingredients = cell_text(row, ing_col)
            out.append({
                'date':            day_date,
                'meal_slot':       slot,
                'dish_freetext':   dish[:200],
                'ingredients_raw': ingredients,
            })
    return out


class Command(BaseCommand):
    help = "Import one Kitchen Coordination biweekly .docx into Menu rows (Mon-Fri only)."

    def add_arguments(self, parser):
        parser.add_argument("docx_path", type=str)
        parser.add_argument("--year", type=int, default=datetime.now().year,
                            help="Year for the date range (default: current year)")
        parser.add_argument("--dry-run", action="store_true")

    def handle(self, *args, **opts):
        path = Path(opts["docx_path"])
        if not path.exists():
            raise CommandError(f"Not found: {path}")

        week1_start, week2_start = parse_date_range(path.name, opts["year"])
        self.stdout.write(f"Week 1 starts: {week1_start}  |  Week 2 starts: {week2_start}")

        doc = docx.Document(str(path))
        if len(doc.tables) < 2:
            raise CommandError(f"Expected 2 tables, found {len(doc.tables)}")

        rows: list[dict] = []
        for week_idx, week_start in enumerate([week1_start, week2_start]):
            table = doc.tables[week_idx]
            fmt = detect_format(table)
            if fmt == 'old':
                # Old format: 3 blocks per table (Mon-Tue, Wed-Thu, Fri-Sat),
                # each 12 rows starting at row 2.
                for block_idx, block_start in enumerate((2, 14, 26)):
                    if block_start + OLD_BLOCK_SIZE > len(table.rows):
                        continue
                    rows.extend(parse_block_old(table, block_start, block_idx, week_start))
            else:
                # New format: blocks at 0, 12, 24
                for block_start in (0, BLOCK_SIZE, BLOCK_SIZE * 2):
                    if block_start + BLOCK_SIZE > len(table.rows):
                        continue
                    rows.extend(parse_block(table, block_start, week_start))

        self.stdout.write(f"Parsed {len(rows)} menu entries")

        if opts["dry_run"]:
            for r in rows:
                self.stdout.write(f"  {r['date']} {r['meal_slot']}: {r['dish_freetext']}")
            return

        created, updated = 0, 0
        for r in rows:
            _, was_created = Menu.objects.update_or_create(
                date=r['date'], meal_slot=r['meal_slot'],
                defaults={
                    'dish_freetext':   r['dish_freetext'],
                    'ingredients_raw': r['ingredients_raw'],
                },
            )
            if was_created:
                created += 1
            else:
                updated += 1
        self.stdout.write(self.style.SUCCESS(f"Menu rows: {created} created, {updated} updated"))
